import os
import re
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
doc.add_heading('VOx Memristor: First-Principles Electro-Thermal Compact Model', 0)
subtitle = doc.add_paragraph('Comprehensive Design, Derivation, and Validation Report')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('Name: _________________ | Course: _________________ | Date: _________________').alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_page_break()

doc.add_heading('Table of Contents', level=1)
doc.add_paragraph('Please update the Table of Contents via Word.')
doc.add_page_break()

with open('report.txt', 'r') as f:
    lines = f.readlines()

figure_count = 1
unplaced_figures = [
    "output/plots/adv_fig3_NDR.png",
    "output/plots/adv_fig5_multi.png",
    "output/plots/real_fig5_multi.png",
    "output/plots/real_fig6_sweeps.png",
    "output/plots/adv_dashboard.png",
    "output/plots/real_fig2_RT.png",
    "output/plots/real_dashboard.png"
]

def add_image(image_path, caption_text):
    global figure_count
    if image_path in unplaced_figures:
        unplaced_figures.remove(image_path)
    if os.path.exists(image_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(image_path, width=Inches(5.5))
        cap = doc.add_paragraph(f"Figure {figure_count}. {caption_text}")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.style = 'Caption'
        figure_count += 1
    else:
        doc.add_paragraph(f"[MISSING IMAGE: {image_path}]")

in_table = False
table_data = []

def render_table(data):
    if not data: return
    table = doc.add_table(rows=1, cols=len(data[0]))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(data[0]):
        hdr_cells[i].text = header
        hdr_cells[i].paragraphs[0].runs[0].bold = True
    for row in data[1:]:
        row_cells = table.add_row().cells
        for i, val in enumerate(row):
            if i < len(row_cells):
                row_cells[i].text = val

current_section = ""

for i, line in enumerate(lines):
    line = line.strip()
    if not line:
        continue
        
    if i < 2:
        continue

    # Tables
    if line.startswith('|') and '|' in line[1:]:
        if not in_table:
            in_table = True
            table_data = []
        parts = [p.strip() for p in line.split('|')[1:-1]]
        if all(re.match(r'^[-: ]+$', p) for p in parts):
            continue
        table_data.append(parts)
        continue
    else:
        if in_table:
            render_table(table_data)
            in_table = False
            table_data = []

    # Headings
    if re.match(r'^\d+\.\s+[A-Z]', line):
        doc.add_heading(line, level=1)
        current_section = line.split()[0]
    elif re.match(r'^\d+\.\d+\.\s+', line):
        doc.add_heading(line, level=2)
        current_section = line.split()[0]
    elif re.match(r'^\d+\.\d+\.\d+\.\s+', line):
        doc.add_heading(line, level=3)
        current_section = line.split()[0]
    elif line.startswith('•') or line.startswith('-'):
        doc.add_paragraph(line[1:].strip(), style='List Bullet')
    elif '=' in line and len(set(line)) < 5:
        continue
    elif 'C_th * (dT/dt)' in line or 'dφ/dt =' in line or 'φ_eq(T) =' in line or 'T_c(φ)' in line or 'R_i(T) =' in line or 'R_m(T) =' in line or 'R_total =' in line or 'τ_φ(T) =' in line or 'G_th(T) =' in line:
        p = doc.add_paragraph(line)
        p.style = 'Normal'
        p.runs[0].font.name = 'Courier New'
        p.paragraph_format.left_indent = Inches(0.5)
    else:
        doc.add_paragraph(line)

    # Insert Images explicitly by text
    if "Below is the dashboard of the fundamental physics in Voltage-Biased mode" in line:
        add_image("output/plots/dashboard.png", "Voltage-biased dashboard (4-panel) - Ideal Model")
    elif "The distinct R-T hysteresis loop correctly captures" in line:
        add_image("output/plots/rt_curve.png", "R-T hysteresis loop - Ideal Model")
    elif "The I-V curve exhibits the classic asymmetric pinched hysteresis" in line:
        add_image("output/plots/iv_curve.png", "I-V pinched hysteresis - Ideal Model")
    elif "The model closely replicates this S-shaped NDR curve:" in line:
        add_image("output/plots/reproduced_fig3.png", "Reproduced Fig 3 (295K Cooling)")
    elif "The simulation captures this amplitude-dependent progression:" in line:
        add_image("output/plots/reproduced_fig5.png", "Reproduced Fig 5 (Current Bias Sweeps)")
    elif "The model captures this sweep-to-sweep memory evolution:" in line:
        add_image("output/plots/reproduced_fig6.png", "Reproduced Fig 6 (Cycle-to-cycle evolution)")

    # Insert images by section if the section just started or just ended, actually better to check section headers
    if "9.4.1" in current_section and "characteristic" in line.lower():
        add_image("output/plots/real_fig2_RT.png", "Real Model R-T characteristic (cf. Figure 2)")
        current_section = ""
    elif "9.4.5" in current_section and "dashboard" in line.lower():
        add_image("output/plots/real_dashboard.png", "Real Model 4-panel dashboard")
        current_section = ""
    elif "10.3.1" in current_section and "ndr" in line.lower():
        add_image("output/plots/adv_fig3_NDR.png", "Advanced Model NDR (cf. Fig 3)")
        current_section = ""
    elif "10.3.2" in current_section and "amplitude" in line.lower():
        add_image("output/plots/adv_fig5_multi.png", "Advanced Model amplitude-sweep (cf. Fig 5c)")
        current_section = ""
    elif "9.4.3" in current_section and "amplitude" in line.lower():
        add_image("output/plots/real_fig5_multi.png", "Real Model multi-amplitude sweep (cf. Figure 5c)")
        current_section = ""
    elif "9.4.4" in current_section and "cycle" in line.lower():
        add_image("output/plots/real_fig6_sweeps.png", "Real Model multi-cycle voltage sweep (cf. Figure 6)")
        current_section = ""
    elif "10.3.3" in current_section and "dashboard" in line.lower():
        add_image("output/plots/adv_dashboard.png", "Advanced Dashboard (3-ODE system, 4-panel)")
        current_section = ""

if in_table:
    render_table(table_data)

if unplaced_figures:
    doc.add_heading('Unplaced Figures — Needs Manual Placement', level=1)
    for fig in unplaced_figures:
        add_image(fig, f"Unplaced Figure: {os.path.basename(fig)}")

# Add References section if not present
ref_text = 'Rana, A.S., del Valle, J. et al., "Symmetric and tunable resistive switching in VO2 memristors," Scientific Reports 10, 3293 (2020), https://doi.org/10.1038/s41598-020-60373-z'
if not any("Rana, A.S." in l for l in lines):
    doc.add_heading('References', level=1)
    doc.add_paragraph(ref_text)

doc.save('VOx_Memristor_Report.docx')
print("Report generated successfully.")
