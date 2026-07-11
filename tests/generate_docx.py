import os
import sys

# ensure `src` is on sys.path
sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), '..', 'src')))

MD_PATH = os.path.abspath(os.path.join(os.path.dirname(__file__), '..', 'src', 'memristor', 'Docs', 'project_documentation.md'))
OUT_DOCX = os.path.abspath(os.path.join(os.path.dirname(__file__), '..', 'ProjectMemristor_Documentation.docx'))
IMAGES = [
    os.path.abspath(os.path.join(os.path.dirname(__file__), '..', f'hp_vi_set1.png')),
    os.path.abspath(os.path.join(os.path.dirname(__file__), '..', f'hp_x_set1.png')),
    os.path.abspath(os.path.join(os.path.dirname(__file__), '..', f'hp_r_set1.png')),
    os.path.abspath(os.path.join(os.path.dirname(__file__), '..', f'hp_vi_set2.png')),
    os.path.abspath(os.path.join(os.path.dirname(__file__), '..', f'hp_x_set2.png')),
    os.path.abspath(os.path.join(os.path.dirname(__file__), '..', f'hp_r_set2.png')),
    os.path.abspath(os.path.join(os.path.dirname(__file__), '..', f'hp_vi_set3.png')),
    os.path.abspath(os.path.join(os.path.dirname(__file__), '..', f'hp_x_set3.png')),
    os.path.abspath(os.path.join(os.path.dirname(__file__), '..', f'hp_r_set3.png')),
]

try:
    from docx import Document
    from docx.shared import Inches
except ImportError:
    print("python-docx is not installed. Install with: pip install python-docx")
    sys.exit(2)

if not os.path.exists(MD_PATH):
    print(f"Markdown source not found: {MD_PATH}")
    sys.exit(1)

with open(MD_PATH, 'r', encoding='utf-8') as f:
    lines = f.read().splitlines()

doc = Document()

doc.add_heading('ProjectMemristor Documentation', level=0)

in_code_block = False
for line in lines:
    if line.strip().startswith('```'):
        in_code_block = not in_code_block
        continue
    if in_code_block:
        doc.add_paragraph(line)
        continue
    if line.startswith('# '):
        doc.add_heading(line[2:].strip(), level=1)
    elif line.startswith('## '):
        doc.add_heading(line[3:].strip(), level=2)
    elif line.startswith('### '):
        doc.add_heading(line[4:].strip(), level=3)
    elif line.strip() == '':
        doc.add_paragraph('')
    else:
        doc.add_paragraph(line)

# Add images with captions
for img_path in IMAGES:
    if os.path.exists(img_path):
        doc.add_page_break()
        doc.add_heading(os.path.basename(img_path), level=2)
        # Insert image, scaled to 6 inches wide if possible
        try:
            doc.add_picture(img_path, width=Inches(6))
        except Exception:
            # fallback: insert without sizing
            doc.add_picture(img_path)
        doc.add_paragraph(f'Figure: {os.path.basename(img_path)}')
    else:
        doc.add_paragraph(f'Image not found: {img_path}')

# Save
try:
    doc.save(OUT_DOCX)
    print(f'Saved DOCX to {OUT_DOCX}')
except Exception as e:
    print('Failed to save DOCX:', e)
    sys.exit(1)
